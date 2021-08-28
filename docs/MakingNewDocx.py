import docx
import os
os.chdir('D:\\Python\\Udemy-AutomateTheBoringStuff\\Excel , Work and PDF Documents\\Reading and Editing Word Document\\DocumentsMade\\Made')
a = docx.Document()
b = a.add_paragraph('Hello This Is A Paragraph. ')
c = a.add_paragraph('This Is Another Paragraph')
d = a.paragraphs[0]
d.add_run(' This Is A New Run')
d.runs[1].bold = True
a.save('Made2.docx')
